#!/usr/bin/env python3
"""
Generate Digital Future NH PAC administrative docs:
  1. Cover letter explaining the PAC (DOCX + PDF)
  2. Wire transfer instructions for TD Bank NH (DOCX + PDF)

Both documents use the brand-book horizontal logo top-left,
Space Grotesk-style headings (Inter fallback for Word), and the
brand color #0A1128.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import subprocess

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / 'docs'
OUT.mkdir(exist_ok=True)
LOGO = ROOT / 'images' / 'logo-horizontal-trans.png'

NAVY = RGBColor(0x0A, 0x11, 0x28)
GRAY = RGBColor(0x4B, 0x55, 0x66)

TODAY = "April 29, 2026"

def add_logo(doc, width_in=2.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(6)

def set_default_font(doc, name='Calibri', size=11):
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.find(qn('w:rFonts')) if rpr is not None else None
    if rfonts is None and rpr is not None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    if rfonts is not None:
        for k in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
            rfonts.set(qn(k), name)

def heading(doc, text, size=18, color=NAVY, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = True
    return p

def para(doc, text, size=11, color=None, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p

def divider(doc, color='0A1128'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def shaded_para(doc, text, fill='F2F4F7', size=10, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.15); pf.right_indent = Inches(0.15)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color
    return p

def margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def footer(doc, text):
    section = doc.sections[0]
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(text)
    r.font.size = Pt(8); r.font.color.rgb = GRAY


# ---------- Cover Letter (one page) ----------
def build_cover_letter():
    doc = Document()
    set_default_font(doc, size=10)
    margins(doc, top=0.5, bottom=0.5, left=0.85, right=0.85)

    add_logo(doc, width_in=2.0)
    divider(doc)

    para(doc, TODAY, color=GRAY, size=10, space_after=8)

    para(doc, "To Whom It May Concern:", bold=True, space_after=6)

    para(doc,
        "Thank you for taking a moment with this letter. Digital Future NH PAC is a "
        "newly registered, non-partisan New Hampshire political action committee. We "
        "raise money, endorse New Hampshire candidates, and contribute to candidates "
        "who defend digital property, financial privacy, and the right of New Hampshire "
        "residents to build and use new technology.",
        size=10, space_after=6)

    para(doc,
        "We are organized under New Hampshire RSA 664 and report contributions to the "
        "New Hampshire Secretary of State at the thresholds the law requires. We accept "
        "contributions in U.S. dollars and in Bitcoin, with crypto contributions valued "
        "at the spot rate at time of receipt and reported the same way as cash.",
        size=10, space_after=6)

    heading(doc, "What we are working on", size=11, space_before=4, space_after=2)
    para(doc,
        "New Hampshire is already the strongest state in the country on digital-asset "
        "policy. It was the first state to enact a Strategic Bitcoin Reserve law (HB 302, "
        "2025, Chapter 4) and the first to grant decentralized autonomous organizations "
        "legal personhood through a permissionless-blockchain registry (RSA 301-B:14, "
        "effective July 2025). Our job is to help keep that ground &mdash; by endorsing "
        "and contributing to candidates from any party who will defend self-custody, "
        "oppose central bank digital currencies, protect the right to mine and run nodes, "
        "recognize smart contracts under state contract law, and support the affordable, "
        "reliable energy supply that compute requires.".replace('&mdash;', '—'),
        size=10, space_after=6)

    heading(doc, "PAC information", size=11, space_before=4, space_after=2)

    info_lines = [
        ("Committee name", "Digital Future NH PAC"),
        ("State of registration", "New Hampshire (RSA 664)"),
        ("EIN", "42-2221221"),
        ("Mailing address", "248 Carley Road, Peterborough, NH 03458"),
        ("Chair", "Christopher Maidment"),
        ("Phone", "540-598-1130 (m)"),
        ("Email", "info@digitalfuturenh.com"),
        ("Website", "digitalfuturenh.com"),
    ]
    table = doc.add_table(rows=len(info_lines), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.7)
    for i, (k, v) in enumerate(info_lines):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.width = Inches(2.0); c1.width = Inches(4.7)
        p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        r0.font.size = Pt(9.5); r0.font.bold = True; r0.font.color.rgb = NAVY
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        r1.font.size = Pt(9.5); r1.font.color.rgb = NAVY

    para(doc,
        "If you need anything further to open or maintain a relationship with the PAC "
        "&mdash; W-9, Secretary of State filing confirmation, EIN letter, or a treasurer "
        "signature card &mdash; please reach out and we will turn it around the same "
        "day.".replace('&mdash;', '—'),
        size=10, space_after=4)

    para(doc, "Sincerely,", size=10, space_after=30)

    # Signature line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1'); top.set(qn('w:color'), '0A1128')
    pBdr.append(top); pPr.append(pBdr)
    r = p.add_run("Christopher Maidment")
    r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Chair, Digital Future NH PAC")
    r.font.size = Pt(9.5); r.font.color.rgb = GRAY

    footer(doc,
        "Paid for by NH Digital Future PAC, 248 Carley Road, Peterborough, NH 03458. Chris Maidment, Chair. Not authorized by any candidate or candidate's committee. Contributions are not tax deductible.")

    out = OUT / 'DigitalFutureNH_Cover_Letter.docx'
    doc.save(out)
    return out


# ---------- Wire Instructions (one page) ----------
def build_wire_instructions():
    doc = Document()
    set_default_font(doc, size=10)
    margins(doc, top=0.5, bottom=0.5, left=0.85, right=0.85)

    add_logo(doc, width_in=2.0)
    divider(doc)

    heading(doc, "Wire Transfer Instructions", size=18, space_before=0, space_after=2)
    para(doc, f"Issued {TODAY}", color=GRAY, size=9, space_after=8)

    heading(doc, "Beneficiary (account holder)", size=11, space_before=4, space_after=2)
    _kv_table(doc, [
        ("Name on account", "Digital Future NH PAC"),
        ("Mailing address", "248 Carley Road, Peterborough, NH 03458"),
        ("EIN", "42-2221221"),
    ])

    heading(doc, "Beneficiary bank", size=11, space_before=8, space_after=2)
    _kv_table(doc, [
        ("Bank name", "TD Bank, N.A."),
        ("Bank location", "Concord, New Hampshire"),
        ("Account number", "9248656542"),
        ("Account type", "Business checking"),
        ("Reference / memo", "Digital Future NH PAC contribution"),
    ])

    heading(doc, "Routing numbers", size=11, space_before=8, space_after=2)
    _kv_table(doc, [
        ("FedWire ABA (incoming wires)", "031101266"),
        ("ACH routing (electronic transfers)", "011400071"),
    ])

    heading(doc, "Notes for senders", size=11, space_before=8, space_after=2)
    bullets = [
        "Include the sender's full name, address, and (above the NH disclosure threshold) occupation and employer in the wire memo.",
        "Anonymous contributions above the NH statutory threshold will be returned.",
        "Corporate-treasury contributions earmarked for specific candidates or bills will not be accepted.",
        "Foreign-national contributions are prohibited under federal law and will be rejected.",
        "After sending, email the wire confirmation to info@digitalfuturenh.com for a contribution receipt.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs:
            r.font.size = Pt(9)

    heading(doc, "Contact", size=11, space_before=8, space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Christopher Maidment, Chair  ·  ")
    r.font.size = Pt(10); r.font.color.rgb = NAVY; r.font.bold = True
    r = p.add_run("540-598-1130 (m)  ·  info@digitalfuturenh.com")
    r.font.size = Pt(10); r.font.color.rgb = GRAY

    footer(doc,
        "Paid for by NH Digital Future PAC, 248 Carley Road, Peterborough, NH 03458. Chris Maidment, Chair. Not authorized by any candidate or candidate's committee. Contributions are not tax deductible.")

    out = OUT / 'DigitalFutureNH_Wire_Instructions.docx'
    doc.save(out)
    return out


def _kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    t.columns[0].width = Inches(2.4)
    t.columns[1].width = Inches(4.1)
    for i, (k, v) in enumerate(rows):
        c0 = t.rows[i].cells[0]; c0.width = Inches(2.4)
        c1 = t.rows[i].cells[1]; c1.width = Inches(4.1)
        p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(k); r0.font.size = Pt(10); r0.font.bold = True; r0.font.color.rgb = NAVY
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(v); r1.font.size = Pt(10); r1.font.color.rgb = NAVY


# ---------- Memorandum of Association (founding charter) ----------
def build_memorandum_of_association():
    doc = Document()
    set_default_font(doc, size=11)
    margins(doc, top=0.6, bottom=0.7, left=1.0, right=1.0)

    add_logo(doc, width_in=1.9)
    divider(doc)

    heading(doc, "Memorandum of Association",
            size=20, space_before=2, space_after=2)
    para(doc, "of NH Digital Future PAC", size=12, color=GRAY, space_after=12)

    para(doc,
        "This Memorandum of Association sets forth the formation, identity, "
        "purpose, and governing principals of NH Digital Future PAC, an "
        "unincorporated political action committee organized under the laws "
        "of the State of New Hampshire.",
        size=11, space_after=10)

    heading(doc, "1. Name", size=12, space_before=8, space_after=2)
    para(doc,
        "The legal name of the committee is NH Digital Future PAC "
        "(the \"Committee\"). The Committee may also be referred to in "
        "branding and public-facing materials as \"Digital Future NH.\"",
        size=11, space_after=8)

    heading(doc, "2. Form of organization", size=12, space_before=6, space_after=2)
    para(doc,
        "The Committee is an unincorporated political committee organized "
        "and registered with the New Hampshire Secretary of State pursuant "
        "to New Hampshire RSA 664 (Political Expenditures and Contributions). "
        "The Committee is not incorporated and is not formed as a business "
        "entity for profit.",
        size=11, space_after=8)

    heading(doc, "3. Federal tax classification", size=12, space_before=6, space_after=2)
    para(doc,
        "The Committee operates as a political organization within the "
        "meaning of Section 527 of the Internal Revenue Code of 1986, as "
        "amended. The Committee has been issued Employer Identification "
        "Number 42-2221221 by the Internal Revenue Service.",
        size=11, space_after=8)

    heading(doc, "4. Registered office and mailing address",
            size=12, space_before=6, space_after=2)
    para(doc,
        "The principal office and mailing address of the Committee is "
        "248 Carley Road, Peterborough, New Hampshire 03458. All official "
        "notices and correspondence may be sent to this address or to "
        "info@digitalfuturenh.com.",
        size=11, space_after=8)

    heading(doc, "5. Purposes", size=12, space_before=6, space_after=2)
    para(doc,
        "The Committee is organized for the following non-partisan public "
        "purposes:",
        size=11, space_after=4)
    for b in [
        "To raise money and accept contributions in compliance with applicable "
        "New Hampshire and federal law;",
        "To endorse and contribute to candidates for public office in New "
        "Hampshire — without regard to political party — who support sound "
        "policy on digital property, blockchain, financial privacy, and the "
        "right of New Hampshire residents to build and use new technology;",
        "To educate New Hampshire voters and policymakers on issues affecting "
        "digital assets, encryption, self-custody, and emerging technology;",
        "To engage in any other lawful activity that advances the foregoing "
        "purposes and is permitted to a New Hampshire political committee "
        "under RSA 664 and applicable federal law.",
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs: r.font.size = Pt(11)

    heading(doc, "6. Powers", size=12, space_before=6, space_after=2)
    para(doc,
        "In furtherance of its purposes, the Committee shall have the power "
        "to: (a) solicit and accept contributions of money, property, and "
        "services in U.S. dollars and in digital assets including Bitcoin; "
        "(b) make contributions to candidate committees and other political "
        "committees as permitted by law; (c) make independent expenditures; "
        "(d) hire and compensate staff and contractors; (e) maintain bank "
        "accounts and merchant payment accounts in its own name; (f) enter "
        "into contracts; (g) own and use intellectual property, including "
        "the trade name \"Digital Future NH\" and associated marks; and "
        "(h) take any other action necessary or convenient to carry out its "
        "lawful purposes.",
        size=11, space_after=8)

    heading(doc, "7. Founding officer", size=12, space_before=6, space_after=2)
    para(doc,
        "The founding officer and Chair of the Committee is Christopher "
        "Maidment. The Chair has authority to act on behalf of the Committee, "
        "including authority to open and operate financial accounts, sign "
        "contracts, accept contributions, retain counsel and contractors, "
        "and file reports required by law. The Chair may delegate ministerial "
        "duties to a Treasurer or other officer as the governing rules of "
        "the Committee may provide.",
        size=11, space_after=8)

    heading(doc, "8. Compliance", size=12, space_before=6, space_after=2)
    para(doc,
        "The Committee shall comply with all applicable provisions of New "
        "Hampshire RSA 664 governing political expenditures and contributions, "
        "including without limitation the registration, recordkeeping, and "
        "disclosure requirements of that chapter. The Committee shall not "
        "knowingly accept contributions from prohibited sources, including "
        "foreign nationals, anonymous contributors above the statutory "
        "threshold, or corporate treasury funds earmarked for specific "
        "candidates.",
        size=11, space_after=8)

    heading(doc, "9. No member financial interest", size=12, space_before=6, space_after=2)
    para(doc,
        "The Committee is non-profit in form. No part of the net assets or "
        "income of the Committee shall inure to the personal benefit of any "
        "officer, founder, contributor, or other private person, except that "
        "the Committee may pay reasonable compensation for services actually "
        "rendered and may make payments and distributions in furtherance of "
        "its lawful purposes.",
        size=11, space_after=8)

    heading(doc, "10. Dissolution", size=12, space_before=6, space_after=2)
    para(doc,
        "Upon dissolution of the Committee, after payment of or provision "
        "for all of the Committee's liabilities, any remaining assets shall "
        "be distributed in accordance with applicable New Hampshire and "
        "federal law to one or more political committees, candidates, or "
        "lawful organizations consistent with the purposes set forth above, "
        "as the Chair may determine.",
        size=11, space_after=8)

    heading(doc, "11. Effective date", size=12, space_before=6, space_after=2)
    para(doc,
        f"This Memorandum of Association is effective as of "
        f"{TODAY}, the date of execution below.",
        size=11, space_after=14)

    para(doc, "Executed by the founding officer:", size=11, space_after=42)

    # Signature
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1'); top.set(qn('w:color'), '0A1128')
    pBdr.append(top); pPr.append(pBdr)
    r = p.add_run("Christopher Maidment, Chair")
    r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
    p = doc.add_paragraph()
    r = p.add_run("NH Digital Future PAC  ·  " + TODAY)
    r.font.size = Pt(10); r.font.color.rgb = GRAY

    footer(doc,
        "Paid for by NH Digital Future PAC, 248 Carley Road, Peterborough, "
        "NH 03458. Chris Maidment, Chair. Not authorized by any candidate or "
        "candidate's committee. Contributions are not tax deductible.")

    out = OUT / 'DigitalFutureNH_Memorandum_of_Association.docx'
    doc.save(out)
    return out


# ---------- Articles of Association (bylaws / operating rules) ----------
def build_articles_of_association():
    doc = Document()
    set_default_font(doc, size=11)
    margins(doc, top=0.6, bottom=0.7, left=1.0, right=1.0)

    add_logo(doc, width_in=1.9)
    divider(doc)

    heading(doc, "Articles of Association",
            size=20, space_before=2, space_after=2)
    para(doc, "of NH Digital Future PAC", size=12, color=GRAY, space_after=4)
    para(doc, f"Adopted {TODAY}", size=10, color=GRAY, space_after=12)

    para(doc,
        "These Articles of Association set forth the internal governing "
        "rules of NH Digital Future PAC (the \"Committee\"), an unincorporated "
        "political committee organized under New Hampshire RSA 664. These "
        "Articles operate together with the Committee's Memorandum of "
        "Association and supersede any prior internal rules.",
        size=11, space_after=12)

    sections = [
        ("Article I — Name and Office",
         "The name of the Committee is NH Digital Future PAC. The principal "
         "office of the Committee is at 248 Carley Road, Peterborough, New "
         "Hampshire 03458, and may be relocated by action of the Chair upon "
         "notice filed with the New Hampshire Secretary of State."),

        ("Article II — Purposes",
         "The Committee is organized to support candidates for public office "
         "in New Hampshire — regardless of political party — who support "
         "sound policy on digital property, blockchain, financial privacy, "
         "and the right of New Hampshire residents to build and use emerging "
         "technology, and to engage in related lawful political activity. The "
         "full statement of purpose is set forth in the Committee's "
         "Memorandum of Association."),

        ("Article III — Members",
         "The Committee has no members. Donors, supporters, volunteers, and "
         "endorsed candidates are not members of the Committee and do not, "
         "by reason of their support, acquire any voting or governance "
         "rights with respect to the Committee."),

        ("Article IV — Officers",
         "The officers of the Committee shall be a Chair and a Treasurer. "
         "The same individual may serve in both offices unless separation "
         "is required by law. Officers serve until they resign, are removed, "
         "or their successor takes office. The initial Chair and Treasurer "
         "is Christopher Maidment."),

        ("Article V — Authority of the Chair",
         "The Chair is the chief officer of the Committee and has full "
         "authority to act on its behalf, including without limitation the "
         "authority to: (a) accept and refuse contributions; (b) make "
         "contributions to candidates and political committees; (c) authorize "
         "expenditures; (d) open, operate, and close financial and merchant "
         "accounts in the Committee's name; (e) sign contracts and "
         "instruments; (f) hire and compensate staff, vendors, counsel, and "
         "advisors; (g) endorse candidates; (h) speak publicly on behalf of "
         "the Committee; and (i) appoint and remove other officers. The "
         "Chair may delegate any of these powers in writing."),

        ("Article VI — Authority of the Treasurer",
         "The Treasurer is responsible for the Committee's financial "
         "recordkeeping and for the timely filing of reports required by "
         "New Hampshire RSA 664 and applicable federal law. The Treasurer "
         "may co-sign disbursements and shall countersign any check or "
         "transfer above a threshold set by the Chair from time to time. "
         "If the offices of Chair and Treasurer are held by the same "
         "individual, the dual-signature requirement is satisfied by that "
         "individual's signature alone."),

        ("Article VII — Decision-making",
         "Decisions of the Committee are made by the Chair. The Chair may, "
         "but is not required to, consult an advisory body of supporters "
         "or counsel. No vote, quorum, or written consent of any other "
         "person is required for any action taken by the Chair within the "
         "scope of the authority set forth in these Articles."),

        ("Article VIII — Banking and payment processing",
         "The Committee shall maintain at least one operating bank account "
         "in its legal name. As of the date of these Articles the Committee "
         "maintains a business checking account at TD Bank, N.A. (Concord, "
         "New Hampshire). The Committee may also maintain merchant "
         "processing accounts (including Stripe, BitPay, or other lawful "
         "processors) for the acceptance of card and digital-asset "
         "contributions. The Chair is the authorized signer and agent for "
         "all such accounts."),

        ("Article IX — Contributions",
         "The Committee accepts contributions in U.S. dollars and in "
         "Bitcoin. Contributions in digital assets shall be valued at the "
         "spot rate at the time of receipt and reported the same way as "
         "cash contributions. The Committee shall not knowingly accept: "
         "(a) contributions from foreign nationals; (b) contributions from "
         "corporate treasuries earmarked for specific candidates or specific "
         "ballot questions; (c) anonymous contributions above the threshold "
         "set by RSA 664; or (d) any contribution prohibited by applicable "
         "law. Improper contributions shall be returned to the donor or "
         "transferred to the United States Treasury as required by law."),

        ("Article X — Books, records, and reports",
         "The Committee shall maintain accurate books and records of all "
         "contributions and expenditures sufficient to comply with RSA 664 "
         "and Section 527 of the Internal Revenue Code. The Committee shall "
         "file with the New Hampshire Secretary of State all reports "
         "required by law and shall make any additional disclosures required "
         "by federal tax law. The fiscal year of the Committee is the "
         "calendar year."),

        ("Article XI — Compliance",
         "The Committee shall comply with all applicable provisions of New "
         "Hampshire RSA 664 and Section 527 of the Internal Revenue Code, "
         "including registration, recordkeeping, disclosure, and reporting "
         "requirements. No officer or agent of the Committee may waive any "
         "such requirement."),

        ("Article XII — Conflicts of interest",
         "Officers shall not use their position to obtain personal benefit "
         "beyond reasonable compensation for services actually rendered. "
         "Any contract or arrangement between the Committee and an officer "
         "(or an entity in which an officer has a material interest) shall "
         "be on arm's-length terms and shall be disclosed in the Committee's "
         "books and records."),

        ("Article XIII — Indemnification",
         "To the fullest extent permitted by law, the Committee shall "
         "indemnify any officer or agent of the Committee against expenses "
         "actually and reasonably incurred in connection with any action, "
         "suit, or proceeding to which such person is made a party by "
         "reason of being an officer or agent of the Committee, except in "
         "matters as to which such person is finally adjudged to have acted "
         "with willful misconduct or in violation of law."),

        ("Article XIV — Amendments",
         "These Articles may be amended at any time by written instrument "
         "signed by the Chair. Amendments take effect on the date stated "
         "in the instrument."),

        ("Article XV — Dissolution",
         "The Committee may be dissolved by written instrument signed by "
         "the Chair. Upon dissolution, after payment of or provision for "
         "all liabilities of the Committee, any remaining assets shall be "
         "distributed in accordance with applicable law to one or more "
         "political committees, candidates, or other lawful organizations "
         "consistent with the purposes of the Committee, as the Chair shall "
         "determine."),
    ]
    for title, body in sections:
        heading(doc, title, size=12, space_before=8, space_after=2)
        para(doc, body, size=11, space_after=4)

    para(doc, " ", size=4, space_after=18)
    para(doc, "Adopted by the Chair of NH Digital Future PAC:", size=11, space_after=42)

    # Signature
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1'); top.set(qn('w:color'), '0A1128')
    pBdr.append(top); pPr.append(pBdr)
    r = p.add_run("Christopher Maidment, Chair")
    r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
    p = doc.add_paragraph()
    r = p.add_run("NH Digital Future PAC  ·  " + TODAY)
    r.font.size = Pt(10); r.font.color.rgb = GRAY

    footer(doc,
        "Paid for by NH Digital Future PAC, 248 Carley Road, Peterborough, "
        "NH 03458. Chris Maidment, Chair. Not authorized by any candidate or "
        "candidate's committee. Contributions are not tax deductible.")

    out = OUT / 'DigitalFutureNH_Articles_of_Association.docx'
    doc.save(out)
    return out


# ---------- Certificate of Authority (one page) ----------
def build_certificate_of_authority():
    doc = Document()
    set_default_font(doc, size=11)
    margins(doc, top=0.6, bottom=0.7, left=1.0, right=1.0)

    add_logo(doc, width_in=1.9)
    divider(doc)

    heading(doc, "Certificate of Authority",
            size=20, space_before=2, space_after=2)
    para(doc, "of NH Digital Future PAC", size=12, color=GRAY, space_after=4)
    para(doc, f"Issued {TODAY}", size=10, color=GRAY, space_after=14)

    para(doc,
        "I, Christopher Maidment, being the duly serving Chair of NH Digital "
        "Future PAC, an unincorporated political committee organized under "
        "the laws of the State of New Hampshire and registered with the New "
        "Hampshire Secretary of State pursuant to RSA 664, do hereby certify "
        "as follows:",
        size=11, space_after=10)

    heading(doc, "1. The Committee", size=12, space_before=4, space_after=2)
    _kv_table(doc, [
        ("Legal name", "NH Digital Future PAC"),
        ("Form of organization", "Unincorporated political committee (NH RSA 664)"),
        ("Federal tax classification", "Section 527 political organization"),
        ("Employer Identification Number", "42-2221221"),
        ("Principal office", "248 Carley Road, Peterborough, NH 03458"),
        ("Email", "info@digitalfuturenh.com"),
        ("Website", "digitalfuturenh.com"),
    ])

    heading(doc, "2. Authorized signer", size=12, space_before=8, space_after=2)
    _kv_table(doc, [
        ("Name", "Christopher Maidment"),
        ("Title", "Chair (and Treasurer)"),
        ("Phone", "540-598-1130 (m)"),
        ("Email", "info@digitalfuturenh.com"),
    ])

    heading(doc, "3. Scope of authority", size=12, space_before=8, space_after=2)
    para(doc,
        "Pursuant to the Memorandum of Association and Articles of "
        "Association of the Committee, the Authorized Signer is fully "
        "authorized, acting alone and without the approval of any other "
        "person, to take the following actions on behalf of the Committee:",
        size=11, space_after=4)
    for b in [
        "Open, maintain, operate, and close bank accounts and merchant "
        "payment processing accounts (including with TD Bank, N.A., Stripe, "
        "BitPay, and any other lawful provider) in the name of the Committee;",
        "Accept, refuse, and refund contributions in U.S. dollars and in "
        "digital assets (including Bitcoin);",
        "Disburse funds, sign checks, initiate wire transfers and ACH "
        "transactions, and authorize card payments;",
        "Execute and deliver applications, agreements, terms of service, "
        "merchant processing agreements, and any related documents on "
        "behalf of the Committee;",
        "Provide identifying information, beneficial-ownership "
        "certifications, and supporting documents to financial "
        "institutions and payment processors;",
        "Make political contributions and independent expenditures, and "
        "publish endorsements, on behalf of the Committee;",
        "File reports required by New Hampshire RSA 664 and Section 527 "
        "of the Internal Revenue Code; and",
        "Take any other action reasonably necessary or convenient to "
        "carry out the lawful purposes of the Committee.",
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs: r.font.size = Pt(10.5)

    heading(doc, "4. Specimen signature", size=12, space_before=8, space_after=2)
    para(doc,
        "The signature appearing below is the genuine signature of the "
        "Authorized Signer named in Section 2 and may be relied upon by "
        "any financial institution, payment processor, counterparty, or "
        "other person doing business with the Committee.",
        size=11, space_after=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1'); top.set(qn('w:color'), '0A1128')
    pBdr.append(top); pPr.append(pBdr)
    r = p.add_run("Christopher Maidment, Chair")
    r.font.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)
    p = doc.add_paragraph()
    r = p.add_run("NH Digital Future PAC  ·  " + TODAY)
    r.font.size = Pt(10); r.font.color.rgb = GRAY

    heading(doc, "5. Effective date and reliance", size=12, space_before=14, space_after=2)
    para(doc,
        f"This Certificate of Authority is effective as of {TODAY} and "
        "remains in full force until revoked in writing by an instrument "
        "signed by the Chair of the Committee and delivered to the "
        "relying party. Any party may rely on this Certificate without "
        "further inquiry.",
        size=11, space_after=8)

    footer(doc,
        "Paid for by NH Digital Future PAC, 248 Carley Road, Peterborough, "
        "NH 03458. Chris Maidment, Chair. Not authorized by any candidate or "
        "candidate's committee. Contributions are not tax deductible.")

    out = OUT / 'DigitalFutureNH_Certificate_of_Authority.docx'
    doc.save(out)
    return out


def docx_to_pdf(docx_path):
    soffice = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    subprocess.run(
        [soffice, '--headless', '--convert-to', 'pdf', '--outdir', str(OUT), str(docx_path)],
        check=True, capture_output=True)
    return docx_path.with_suffix('.pdf')


if __name__ == '__main__':
    cl = build_cover_letter()
    wi = build_wire_instructions()
    moa = build_memorandum_of_association()
    aoa = build_articles_of_association()
    coa = build_certificate_of_authority()
    print('DOCX written:')
    for f in (cl, wi, moa, aoa, coa):
        print(' ', f)
    print('Converting to PDF...')
    for d in (cl, wi, moa, aoa, coa):
        p = docx_to_pdf(d)
        print('  ->', p)
